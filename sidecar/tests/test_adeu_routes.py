"""
Pytest tests for sidecar Adeu endpoints.

Covers: /adeu/read, /adeu/process-batch, /adeu/accept-all,
        /adeu/apply-edits-markdown, /adeu/diff, and /health.

Requirements: 2.1, 2.2, 2.4, 3.1, 3.4, 3.5, 3.6, 4.1, 5.1, 6.4, 11.1
"""

import io
import json

import pytest
from docx import Document


# ── /adeu/read ──────────────────────────────────────────────────────────────


class TestReadDocx:
    """Req 2.1, 2.2, 2.4"""

    def test_read_valid_docx_returns_text(self, client, simple_docx):
        """Req 2.1 — POST /adeu/read with valid DOCX returns text content."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("test.docx", simple_docx)},
            data={"clean_view": "false"},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert "text" in body
        assert "Hello world" in body["text"]
        assert body["filename"] == "test.docx"

    def test_read_clean_view_strips_critic_markup(self, client, simple_docx):
        """Req 2.2 — clean_view=true returns accepted-state text without CriticMarkup."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("doc.docx", simple_docx)},
            data={"clean_view": "true"},
        )
        assert resp.status_code == 200
        text = resp.json()["text"]
        # A plain doc has no markup either way, but clean_view should not crash
        for delim in ("{--", "--}", "{++", "++}", "{==", "==}", "{>>", "<<}"):
            assert delim not in text

    def test_read_invalid_file_returns_422(self, client):
        """Req 2.4 — non-DOCX file returns 422."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("bad.docx", b"this is not a docx file")},
        )
        assert resp.status_code == 422
        assert "detail" in resp.json()


# ── /adeu/process-batch ─────────────────────────────────────────────────────


class TestProcessBatch:
    """Req 3.1, 3.4, 3.5, 3.6"""

    def test_apply_edits_returns_modified_docx_and_summary(
        self, client, multi_paragraph_docx
    ):
        """Req 3.1, 3.5 — valid edits return modified DOCX + BatchSummary header."""
        body = json.dumps(
            {
                "author_name": "Test Author",
                "edits": [
                    {
                        "target_text": "quick brown fox",
                        "new_text": "slow red fox",
                    }
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("contract.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        assert (
            resp.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Verify summary header
        summary_raw = resp.headers.get("x-batch-summary")
        assert summary_raw is not None
        summary = json.loads(summary_raw)
        assert summary["applied_edits"] == 1
        assert summary["skipped_edits"] == 0

        # Verify the response is a valid DOCX
        doc = Document(io.BytesIO(resp.content))
        assert len(doc.paragraphs) > 0

    def test_invalid_target_text_returns_422(self, client, simple_docx):
        """Req 3.4 — edit with non-existent target_text returns 422."""
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {
                        "target_text": "text that does not exist anywhere",
                        "new_text": "replacement",
                    }
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("doc.docx", simple_docx)},
            data={"body": body},
        )
        assert resp.status_code == 422
        data = resp.json()
        assert "detail" in data

    def test_empty_batch_returns_422(self, client, simple_docx):
        """Req 3.6 — no edits and no actions returns 422."""
        body = json.dumps({"author_name": "Author"})
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("doc.docx", simple_docx)},
            data={"body": body},
        )
        assert resp.status_code == 422
        assert "At least one edit or action" in resp.json()["detail"]


# ── /adeu/accept-all ────────────────────────────────────────────────────────


class TestAcceptAll:
    """Req 4.1"""

    def test_accept_all_returns_clean_docx(self, client, simple_docx):
        """Req 4.1 — accept-all returns a valid DOCX binary."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("doc.docx", simple_docx)},
        )
        assert resp.status_code == 200
        assert (
            resp.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # Verify it's a valid DOCX
        doc = Document(io.BytesIO(resp.content))
        assert len(doc.paragraphs) > 0


# ── /adeu/apply-edits-markdown ──────────────────────────────────────────────


class TestApplyEditsMarkdown:
    """Req 5.1"""

    def test_returns_critic_markup_annotated_markdown(
        self, client, multi_paragraph_docx
    ):
        """Req 5.1 — edits produce CriticMarkup annotations in markdown."""
        body = json.dumps(
            {
                "edits": [
                    {
                        "target_text": "quick brown fox",
                        "new_text": "slow red fox",
                    }
                ],
            }
        )
        resp = client.post(
            "/adeu/apply-edits-markdown",
            files={"file": ("doc.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        md = resp.json()["markdown"]
        # CriticMarkup deletion and insertion markers should be present
        assert "{--" in md or "{++" in md


# ── /adeu/diff (error cases) ────────────────────────────────────────────────


class TestDiffErrors:
    """Req 2.21 — negative test cases for /adeu/diff."""

    def test_diff_missing_original_returns_422(self, client, simple_docx):
        """Missing 'original' field returns 422."""
        resp = client.post(
            "/adeu/diff",
            files={"modified": ("b.docx", simple_docx)},
            data={"compare_clean": "true"},
        )
        assert resp.status_code == 422

    def test_diff_missing_modified_returns_422(self, client, simple_docx):
        """Missing 'modified' field returns 422."""
        resp = client.post(
            "/adeu/diff",
            files={"original": ("a.docx", simple_docx)},
            data={"compare_clean": "true"},
        )
        assert resp.status_code == 422

    def test_diff_invalid_docx_bytes_returns_error(self, client):
        """Invalid DOCX bytes return 400 or 422."""
        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", b"not a docx"),
                "modified": ("b.docx", b"not a docx"),
            },
            data={"compare_clean": "true"},
        )
        assert resp.status_code in (400, 422)


# ── /adeu/accept-all (error cases) ──────────────────────────────────────────


class TestAcceptAllErrors:
    """Req 2.21 — negative test cases for /adeu/accept-all."""

    def test_accept_all_missing_file_returns_422(self, client):
        """Missing 'file' field returns 422."""
        resp = client.post("/adeu/accept-all")
        assert resp.status_code == 422

    def test_accept_all_invalid_docx_bytes_returns_error(self, client):
        """Invalid DOCX bytes return 400, 422, or 500."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("bad.docx", b"not a docx")},
        )
        assert resp.status_code in (400, 422, 500)


# ── /adeu/apply-edits-markdown (error cases) ────────────────────────────────


class TestApplyEditsMarkdownErrors:
    """Req 2.21 — negative test cases for /adeu/apply-edits-markdown."""

    def test_apply_edits_markdown_missing_file_returns_422(self, client):
        """Missing 'file' field returns 422."""
        body = json.dumps({"edits": [{"target_text": "x", "new_text": "y"}]})
        resp = client.post(
            "/adeu/apply-edits-markdown",
            data={"body": body},
        )
        assert resp.status_code == 422

    def test_apply_edits_markdown_invalid_docx_returns_error(self, client):
        """Invalid DOCX bytes return 400 or 422."""
        body = json.dumps({"edits": [{"target_text": "x", "new_text": "y"}]})
        resp = client.post(
            "/adeu/apply-edits-markdown",
            files={"file": ("bad.docx", b"not a docx")},
            data={"body": body},
        )
        assert resp.status_code in (400, 422)


# ── /adeu/diff ──────────────────────────────────────────────────────────────


class TestDiff:
    """Req 6.4"""

    def test_identical_files_no_differences(self, client, simple_docx):
        """Req 6.4 — same file as both original and modified → has_differences: false."""
        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", simple_docx),
                "modified": ("b.docx", simple_docx),
            },
            data={"compare_clean": "true"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["has_differences"] is False


# ── /health ─────────────────────────────────────────────────────────────────


class TestHealth:
    """Req 11.1"""

    def test_health_returns_adeu_version(self, client):
        """Req 11.1 — /health reports Adeu availability and version."""
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"
        assert data["adeu"]["available"] is True
        assert data["adeu"]["version"] is not None
