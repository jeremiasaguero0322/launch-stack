"""Shared fixtures for sidecar Adeu route tests.

We build a lightweight FastAPI app with only the adeu router and health
endpoint so that tests don't require ML dependencies (torch, etc.).
"""

import io
import os

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.routes.adeu import router as adeu_router

# Test API key used for all authenticated requests
TEST_API_KEY = "test-sidecar-api-key"


def _build_test_app() -> FastAPI:
    """Minimal app containing only the adeu routes + health check."""
    app = FastAPI()
    app.include_router(adeu_router)

    @app.get("/health")
    async def health():
        try:
            from adeu import __version__ as adeu_version
            adeu_status = {"available": True, "version": adeu_version}
        except ImportError:
            adeu_status = {"available": False, "version": None}
        status = "ok" if adeu_status["available"] else "degraded"
        return {"status": status, "adeu": adeu_status}

    return app


@pytest.fixture(autouse=True)
def set_api_key_env(monkeypatch):
    """Set SIDECAR_API_KEY for all tests so auth middleware passes."""
    monkeypatch.setenv("SIDECAR_API_KEY", TEST_API_KEY)


@pytest.fixture
def client():
    """FastAPI test client with adeu routes only, pre-configured with API key header."""
    test_client = TestClient(_build_test_app())
    # Inject the API key header into all requests by default
    test_client.headers.update({"X-API-Key": TEST_API_KEY})
    return test_client


@pytest.fixture
def simple_docx() -> bytes:
    """A minimal valid DOCX with known text content."""
    doc = Document()
    doc.add_paragraph("Hello world. This is a test document.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def multi_paragraph_docx() -> bytes:
    """A DOCX with multiple paragraphs for edit testing."""
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("Pack my box with five dozen liquor jugs.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
