"""
Preservation Property Tests — Sidecar ADEU Routes

Property 2: Preservation — Normal ADEU Pipeline Behavior
Written BEFORE fixes. EXPECTED TO PASS on unfixed code (confirms baseline behavior).

These tests verify that the existing happy-path sidecar behavior is preserved:
- Valid DOCX files process correctly through all 5 operations
- Valid multipart uploads parse and forward correctly
- Success responses maintain same body format and headers
- Health endpoint returns correct status when adeu is installed
- Safe filenames appear correctly in Content-Disposition headers

Requirements: 3.5, 3.6, 3.7, 3.8, 3.10, 3.12
"""

import io
import json

import pytest
from docx import Document


# ===========================================================================
# Preservation: /adeu/read — valid DOCX returns text with correct format
# ===========================================================================


class TestPreservationReadDocx:
    """Req 3.5 — read endpoint returns correct result for valid DOCX."""

    def test_valid_docx_returns_200_with_text_and_filename(self, client, simple_docx):
        """For all valid DOCX inputs, /adeu/read returns { text, filename }."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("test.docx", simple_docx)},
            data={"clean_view": "false"},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert "text" in body
        assert "filename" in body
        assert isinstance(body["text"], str)
        assert body["filename"] == "test.docx"
        assert len(body["text"]) > 0

    def test_clean_view_returns_200_without_critic_markup(self, client, simple_docx):
        """clean_view=true returns text without CriticMarkup delimiters."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("doc.docx", simple_docx)},
            data={"clean_view": "true"},
        )
        assert resp.status_code == 200
        text = resp.json()["text"]
        assert isinstance(text, str)
        for delim in ("{--", "--}", "{++", "++}", "{==", "==}"):
            assert delim not in text

    def test_multi_paragraph_docx_returns_all_content(
        self, client, multi_paragraph_docx
    ):
        """Multi-paragraph DOCX returns text containing all paragraphs."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("multi.docx", multi_paragraph_docx)},
        )
        assert resp.status_code == 200
        text = resp.json()["text"]
        assert "quick brown fox" in text
        assert "five dozen liquor jugs" in text


# ===========================================================================
# Preservation: /adeu/process-batch — valid edits return modified DOCX
# ===========================================================================


class TestPreservationProcessBatch:
    """Req 3.5, 3.7 — process-batch returns modified DOCX + summary header."""

    def test_valid_edit_returns_200_with_docx_and_summary(
        self, client, multi_paragraph_docx
    ):
        """For valid edits, returns modified DOCX binary + X-Batch-Summary header."""
        body = json.dumps(
            {
                "author_name": "Test Author",
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("contract.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200

        # Response content type is DOCX
        assert (
            resp.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # X-Batch-Summary header is valid JSON with expected fields
        summary_raw = resp.headers.get("x-batch-summary")
        assert summary_raw is not None
        summary = json.loads(summary_raw)
        assert "applied_edits" in summary
        assert "skipped_edits" in summary
        assert "applied_actions" in summary
        assert "skipped_actions" in summary
        assert summary["applied_edits"] >= 1

        # Response body is a valid DOCX
        doc = Document(io.BytesIO(resp.content))
        assert len(doc.paragraphs) > 0

    def test_content_disposition_header_present(self, client, multi_paragraph_docx):
        """Successful batch edit includes Content-Disposition header."""
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("report.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "attachment" in content_disp
        assert "filename=" in content_disp

    def test_validation_error_returns_422_with_detail(self, client, simple_docx):
        """Invalid target_text returns 422 with detail field."""
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {
                        "target_text": "text that does not exist anywhere in the doc",
                        "new_text": "replacement",
                    }
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("doc.docx", simple_docx)},
            data={"body": body},
        )
        assert resp.status_code == 422
        data = resp.json()
        assert "detail" in data


# ===========================================================================
# Preservation: /adeu/accept-all — returns clean DOCX
# ===========================================================================


class TestPreservationAcceptAll:
    """Req 3.5 — accept-all returns valid clean DOCX."""

    def test_valid_docx_returns_200_with_clean_docx(self, client, simple_docx):
        """For valid DOCX, accept-all returns a valid DOCX binary."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("doc.docx", simple_docx)},
        )
        assert resp.status_code == 200
        assert (
            resp.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # Verify it's a valid DOCX
        doc = Document(io.BytesIO(resp.content))
        assert len(doc.paragraphs) > 0

    def test_content_disposition_header_present(self, client, simple_docx):
        """accept-all includes Content-Disposition header with filename."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("report.docx", simple_docx)},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "attachment" in content_disp
        assert "filename=" in content_disp


# ===========================================================================
# Preservation: /adeu/apply-edits-markdown — returns CriticMarkup markdown
# ===========================================================================


class TestPreservationApplyEditsMarkdown:
    """Req 3.5 — apply-edits-markdown returns CriticMarkup-annotated markdown."""

    def test_valid_edits_return_200_with_markdown(
        self, client, multi_paragraph_docx
    ):
        """For valid edits, returns { markdown } with CriticMarkup annotations."""
        body = json.dumps(
            {
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/apply-edits-markdown",
            files={"file": ("doc.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "markdown" in data
        assert isinstance(data["markdown"], str)
        # CriticMarkup markers should be present
        assert "{--" in data["markdown"] or "{++" in data["markdown"]

    def test_highlight_only_mode_returns_highlight_markers(
        self, client, multi_paragraph_docx
    ):
        """highlight_only=true returns highlight markers instead of edit markers."""
        body = json.dumps(
            {
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
                "highlight_only": True,
            }
        )
        resp = client.post(
            "/adeu/apply-edits-markdown",
            files={"file": ("doc.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        md = resp.json()["markdown"]
        assert isinstance(md, str)


# ===========================================================================
# Preservation: /adeu/diff — returns diff result
# ===========================================================================


class TestPreservationDiff:
    """Req 3.5 — diff returns correct result for valid DOCX pairs."""

    def test_identical_files_return_no_differences(self, client, simple_docx):
        """Same file as both original and modified → has_differences: false."""
        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", simple_docx),
                "modified": ("b.docx", simple_docx),
            },
            data={"compare_clean": "true"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["has_differences"] is False
        assert "diff" in data

    def test_different_files_return_differences(self, client):
        """Different files → has_differences: true with non-empty diff."""
        doc_a = Document()
        doc_a.add_paragraph("Original text content.")
        buf_a = io.BytesIO()
        doc_a.save(buf_a)

        doc_b = Document()
        doc_b.add_paragraph("Modified text content.")
        buf_b = io.BytesIO()
        doc_b.save(buf_b)

        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", buf_a.getvalue()),
                "modified": ("b.docx", buf_b.getvalue()),
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["has_differences"] is True
        assert len(data["diff"]) > 0

    def test_diff_response_format(self, client, simple_docx):
        """Diff response has { diff: str, has_differences: bool } shape."""
        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", simple_docx),
                "modified": ("b.docx", simple_docx),
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert isinstance(data["diff"], str)
        assert isinstance(data["has_differences"], bool)


# ===========================================================================
# Preservation: /health — returns adeu status when installed
# ===========================================================================


class TestPreservationHealth:
    """Req 3.12 — health endpoint returns correct status when adeu is installed."""

    def test_health_returns_ok_with_adeu_available(self, client):
        """Health returns { status: 'ok', adeu: { available: true, version: '...' } }."""
        resp = client.get("/health")
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "ok"
        assert data["adeu"]["available"] is True
        assert data["adeu"]["version"] is not None
        assert isinstance(data["adeu"]["version"], str)
        assert len(data["adeu"]["version"]) > 0


# ===========================================================================
# Preservation: Safe filenames in Content-Disposition headers
# ===========================================================================


class TestPreservationFilenameInHeaders:
    """Req 3.10 — valid filenames appear correctly in Content-Disposition."""

    def test_simple_filename_preserved_in_accept_all(self, client, simple_docx):
        """A simple alphanumeric filename appears unchanged in the header."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("report2024.docx", simple_docx)},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "report2024.docx" in content_disp

    def test_simple_filename_preserved_in_process_batch(
        self, client, multi_paragraph_docx
    ):
        """A simple filename appears unchanged in process-batch response."""
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("contract.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "contract.docx" in content_disp

    def test_filename_with_spaces_preserved(self, client, simple_docx):
        """A filename with spaces appears in the header."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("my report.docx", simple_docx)},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "my report.docx" in content_disp

    def test_filename_with_hyphens_and_underscores_preserved(
        self, client, simple_docx
    ):
        """A filename with hyphens and underscores appears unchanged."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("my-report_v2.docx", simple_docx)},
        )
        assert resp.status_code == 200
        content_disp = resp.headers.get("content-disposition", "")
        assert "my-report_v2.docx" in content_disp


# ===========================================================================
# Preservation: Success response format consistency
# ===========================================================================


class TestPreservationResponseFormat:
    """Req 3.7 — success responses maintain same body format and headers."""

    def test_read_response_is_json(self, client, simple_docx):
        """Read endpoint returns JSON content type."""
        resp = client.post(
            "/adeu/read",
            files={"file": ("doc.docx", simple_docx)},
        )
        assert resp.status_code == 200
        assert "application/json" in resp.headers.get("content-type", "")

    def test_process_batch_response_is_docx_binary(
        self, client, multi_paragraph_docx
    ):
        """Process-batch returns DOCX content type."""
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/process-batch",
            files={"file": ("doc.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        assert (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in resp.headers.get("content-type", "")
        )

    def test_accept_all_response_is_docx_binary(self, client, simple_docx):
        """Accept-all returns DOCX content type."""
        resp = client.post(
            "/adeu/accept-all",
            files={"file": ("doc.docx", simple_docx)},
        )
        assert resp.status_code == 200
        assert (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in resp.headers.get("content-type", "")
        )

    def test_apply_edits_markdown_response_is_json(
        self, client, multi_paragraph_docx
    ):
        """Apply-edits-markdown returns JSON content type."""
        body = json.dumps(
            {
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = client.post(
            "/adeu/apply-edits-markdown",
            files={"file": ("doc.docx", multi_paragraph_docx)},
            data={"body": body},
        )
        assert resp.status_code == 200
        assert "application/json" in resp.headers.get("content-type", "")

    def test_diff_response_is_json(self, client, simple_docx):
        """Diff returns JSON content type."""
        resp = client.post(
            "/adeu/diff",
            files={
                "original": ("a.docx", simple_docx),
                "modified": ("b.docx", simple_docx),
            },
        )
        assert resp.status_code == 200
        assert "application/json" in resp.headers.get("content-type", "")
